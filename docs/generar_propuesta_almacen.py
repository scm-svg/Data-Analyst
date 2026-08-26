#!/usr/bin/env python3
"""Genera la propuesta corregida de capacidad del Almacén Fábrica."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

NAVY = RGBColor(0x1B, 0x3A, 0x4B)
ACCENT = RGBColor(0x1F, 0x6F, 0x8B)
DARK = RGBColor(0x2B, 0x2B, 0x2B)
MUTED = RGBColor(0x5C, 0x67, 0x70)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
ROW_ALT = "F4F7F8"
HEADER_BG = "1B3A4B"
GREEN_BG = "E5F4EA"
AMBER_BG = "FFF4D6"
BLUE_BG = "E8F1F5"
RED_BG = "FDECEC"
REC_BG = "E8F6EF"


def fmt(n: float, d: int = 2) -> str:
    s = f"{n:,.{d}f}"
    return s.replace(",", "X").replace(".", ",").replace("X", ".")


def set_run_font(run, name="Calibri", size=11, bold=False, italic=False, color=DARK):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def shade_cell(cell, hex_color: str) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    for child in list(tc_pr):
        if child.tag == qn("w:shd"):
            tc_pr.remove(child)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def set_cell_borders(cell, color="C5D0D4", sz="4") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    for child in list(tc_pr):
        if child.tag == qn("w:tcBorders"):
            tc_pr.remove(child)
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:color"), color)
        el.set(qn("w:space"), "0")
        borders.append(el)
    tc_pr.append(borders)


def set_cell_margins(cell, top=60, bottom=60, left=80, right=80) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for name, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = OxmlElement(f"w:{name}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    tc_pr.append(tc_mar)


def clear_cell(cell) -> None:
    for p in cell.paragraphs:
        p.clear()


def write_cell(
    cell,
    text,
    *,
    bold=False,
    size=10,
    color=DARK,
    align="left",
    fill=None,
    italic=False,
) -> None:
    clear_cell(cell)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    set_cell_borders(cell)
    set_cell_margins(cell)
    if fill:
        shade_cell(cell, fill)


def prevent_row_split(row) -> None:
    tr = row._tr
    tr_pr = tr.get_or_add_trPr()
    cant = OxmlElement("w:cantSplit")
    tr_pr.append(cant)


def set_table_widths(table, widths_cm) -> None:
    table.autofit = False
    table.allow_autofit = False
    total = sum(widths_cm)
    table.width = Cm(total)
    for row in table.rows:
        for i, w in enumerate(widths_cm):
            row.cells[i].width = Cm(w)


def is_hex_color(val) -> bool:
    return isinstance(val, str) and len(val) == 6 and all(c in "0123456789ABCDEFabcdef" for c in val)


def normalize_row(row, n_cols):
    items = list(row)
    row_bold = False
    row_fill = None
    if items and isinstance(items[-1], bool):
        row_bold = items.pop()
    if items and is_hex_color(items[-1]):
        row_fill = items.pop()
    if len(items) == 1 and isinstance(items[0], tuple) and len(items[0]) == n_cols and all(
        not is_hex_color(x) for x in items[0]
    ):
        items = list(items[0])
    if len(items) != n_cols:
        raise ValueError(f"Fila con {len(items)} celdas, se esperaban {n_cols}: {items!r}")
    return items, row_fill, row_bold


def add_table(doc, headers, rows, col_widths, header_fill=HEADER_BG):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        write_cell(
            table.rows[0].cells[i],
            h,
            bold=True,
            size=9.5,
            color=WHITE,
            align="center",
            fill=header_fill,
        )
    prevent_row_split(table.rows[0])
    n_cols = len(headers)
    for r_i, row in enumerate(rows):
        items, row_fill, row_bold = normalize_row(row, n_cols)
        default_fill = row_fill or (ROW_ALT if r_i % 2 else "FFFFFF")
        for c_i, val in enumerate(items):
            cell_fill = default_fill
            bold = row_bold or (c_i == 0)
            text = val
            if isinstance(val, tuple):
                text = val[0]
                if len(val) > 1 and val[1]:
                    cell_fill = val[1]
                if len(val) > 2:
                    bold = val[2]
            write_cell(
                table.rows[r_i + 1].cells[c_i],
                str(text),
                bold=bold,
                size=9.5,
                color=DARK,
                align="left" if c_i == 0 else "center",
                fill=cell_fill,
            )
        prevent_row_split(table.rows[r_i + 1])
    set_table_widths(table, col_widths)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)
    spacer.paragraph_format.space_before = Pt(2)
    return table


def add_para(
    doc,
    text,
    *,
    size=11,
    bold=False,
    italic=False,
    color=DARK,
    space_after=8,
    space_before=0,
    align="left",
):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.15
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    return p


def add_rich(doc, parts, *, space_after=8, size=11, justify=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for text, kwargs in parts:
        run = p.add_run(text)
        set_run_font(run, size=kwargs.get("size", size), **{k: v for k, v in kwargs.items() if k != "size"})
    return p


def add_heading_custom(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.1
    if level == 1:
        run = p.add_run(text.upper())
        set_run_font(run, size=14, bold=True, color=NAVY)
        # underline bar via bottom border on paragraph
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), "1F6F8B")
        pBdr.append(bottom)
        pPr.append(pBdr)
    else:
        run = p.add_run(text)
        set_run_font(run, size=12, bold=True, color=ACCENT)
    return p


def add_bullet(doc, text, *, bold_lead=None, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.12
    p.paragraph_format.left_indent = Cm(1.0 + 0.5 * level)
    if bold_lead:
        r1 = p.add_run(bold_lead)
        set_run_font(r1, size=11, bold=True, color=NAVY)
        r2 = p.add_run(text)
        set_run_font(r2, size=11, color=DARK)
    else:
        r = p.add_run(text)
        set_run_font(r, size=11, color=DARK)
    return p


def add_callout(doc, title, body, fill):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    clear_cell(cell)
    shade_cell(cell, fill)
    set_cell_borders(cell, color="D0D8DC", sz="4")
    set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_after = Pt(4)
    r1 = p1.add_run(title)
    set_run_font(r1, size=10.5, bold=True, color=NAVY)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.12
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.5, color=DARK)
    set_table_widths(table, [16.5])
    doc.add_paragraph()


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = DARK
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    pf = normal.paragraph_format
    pf.line_spacing = 1.15
    pf.space_after = Pt(8)


def build() -> Path:
    doc = Document()
    configure_styles(doc)

    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = hp.add_run("Almacén Fábrica  ·  Propuesta de capacidad y estructura  ·  Versión corregida")
        set_run_font(r, size=8.5, color=MUTED)
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        fr = fp.add_run("Confidencial — uso interno    |    Página ")
        set_run_font(fr, size=8.5, color=MUTED)
        add_page_number(fp)
        fr2 = fp.add_run("")
        set_run_font(fr2, size=8.5, color=MUTED)

    # ===== PORTADA / BLOQUE DE APERTURA =====
    add_para(doc, "DOCUMENTO PARA DIRECCIÓN  ·  OPERACIONES / LOGÍSTICA  ·  AGOSTO 2026", size=10, bold=True, color=ACCENT, space_after=4)
    add_para(
        doc,
        "Propuesta de capacidad y reestructuración del Team Almacén Fábrica",
        size=22,
        bold=True,
        color=NAVY,
        space_after=6,
    )
    add_para(
        doc,
        "Versión corregida: números revisados, supuestos explícitos y una petición clara. Reemplaza el borrador de 8 páginas basado en OEE / “3 operadores exactos”.",
        size=12,
        italic=True,
        color=MUTED,
        space_after=10,
    )

    meta = doc.add_table(rows=2, cols=4)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    labels = ["Área", "Plantilla de partida", "Horizonte", "Estado del documento"]
    values = ["Almacén Fábrica", "9 almacenistas (a confirmar)", "12 meses + proyecto Refresh", "Listo para decisión, con datos en ámbar"]
    for i, (lab, val) in enumerate(zip(labels, values)):
        write_cell(meta.rows[0].cells[i], lab, bold=True, size=8.5, color=WHITE, align="center", fill=HEADER_BG)
        write_cell(meta.rows[1].cells[i], val, size=9.5, color=DARK, align="center", fill=BLUE_BG)
    set_table_widths(meta, [4.125, 4.125, 4.125, 4.125])
    doc.add_paragraph()

    add_callout(
        doc,
        "Qué se pide firmar (en una frase)",
        "Reorganizar de inmediato el almacén en 4 células con las 9 personas actuales; abrir 1 vacante permanente; contratar 2 refuerzos temporales (4 a 6 meses) solo para el Proyecto Refresh; y no convertir un pico de esta semana ni un proyecto con fecha de fin en 3 cargos fijos.",
        REC_BG,
    )

    # ===== 1. RESUMEN EJECUTIVO =====
    add_heading_custom(doc, "1. Resumen ejecutivo", 1)
    add_rich(
        doc,
        [
            ("El almacén no está “al 99,6 % de su techo estructural”. ", {"bold": True}),
            (
                "Ese porcentaje sale de restar vacaciones de una persona, dos bajas de salud al 30 % y una curva de aprendizaje, y luego comparar esa foto de la semana contra las horas ya asignadas en el turno. Las sumas del borrador están bien; el modelo no.",
                {},
            ),
        ],
    )
    add_rich(
        doc,
        [
            ("Lo que sí es cierto: ", {"bold": True}),
            (
                "hay tareas de control (conteos, 5S, devoluciones) y un proyecto de 10.000 piezas (Refresh) que hoy no tienen dueño ni horas protegidas. También hay un plan de 1 tienda nueva cada 3 meses que, si está aprobado, sí aumenta la carga de reabastecimiento. Eso merece una respuesta. No merece inflar la plantilla permanente con pérdidas temporales.",
                {},
            ),
        ],
    )

    add_heading_custom(doc, "La decisión, en números", 2)
    add_table(
        doc,
        ["Magnitud", "Horas-hombre / semana", "Equivalente"],
        [
            ["Jornada neta informada (9 personas × 8,25 h × 5 días)", "371,25 HH", "9,00 FTE teóricos"],
            ["Capacidad estructural (vacaciones y ausentismo anualizados)", "335,84 HH", "8,14 FTE efectivos"],
            ["Carga del día a día (horas hoy asignadas, no estándar de ingeniería)", "287,75 HH", "6,98 FTE"],
            ["Tareas omitidas sostenibles (conteos + 5S + devoluciones)", "65,00 HH", "1,58 FTE"],
            [("Proyecto Refresh (10.000 piezas) — no es plantilla fija", "40,00 HH", "0,97 FTE por ~3 meses"), AMBER_BG, True],
            ["Demanda sostenible (día a día + omitidas, sin Refresh)", "352,75 HH", "8,55 FTE"],
            [("Brecha estructural hoy, sin Refresh", "+16,91 HH", "0,41 FTE — no 2,5"), GREEN_BG, True],
            [("Brecha si se trata Refresh como permanente (error del borrador)", "+56,91 a +104 HH", "1,4 a 2,5 FTE"), AMBER_BG, True],
        ],
        [8.4, 4.1, 4.0],
    )

    add_para(doc, "Petición concreta a Dirección", size=12, bold=True, color=NAVY, space_after=6)
    add_bullet(doc, " aprobar el mapa de 4 células y las reglas de vale / acceso. No requiere esperar contrataciones.", bold_lead="Hoy (costo cero de headcount):")
    add_bullet(doc, " 1 Operador Integral de Almacén, con prioridad en Materias Primas (curva de aprendizaje y control de entrega a Producción).", bold_lead="Vacante permanente:")
    add_bullet(doc, " 2 auxiliares o almacenistas por 4–6 meses, con meta semanal de piezas y fecha de fin. Al cerrar Refresh, salen o se evalúan contra aperturas reales.", bold_lead="Refuerzo temporal:")
    add_bullet(doc, " al confirmar la apertura de la tienda 10 (no antes), revisar una segunda permanente en Producto Terminado. No se pide ahora “por si acaso”.", bold_lead="Gatillo de crecimiento:")
    add_bullet(doc, " en 15 días, cerrar si la nómina es 9 u 10; qué hacen hoy jueves y viernes las células de PT; OTIF / ERI actuales; y el valor recuperable de Refresh.", bold_lead="Condición:")

    add_callout(
        doc,
        "Por qué no se piden 3 fijos “exactos”",
        "Con el factor de 77,7 % que usaba el borrador, 3 contrataciones aportan ~96 HH reales: no cubren ni el déficit de 104 HH que el mismo texto calculaba, y menos el año de aperturas. Pedir 3 como número exacto no cierra ni con las propias premisas. La petición de este documento sí cuadra: 1 FTE estructural (~0,4 de brecha run-rate + holgura de onboarding y MP) y 2 temporales para las ~40 HH de Refresh.",
        AMBER_BG,
    )

    # ===== 2. CORRECCIONES =====
    add_heading_custom(doc, "2. Qué se corrigió del borrador anterior", 1)
    add_para(
        doc,
        "El borrador estaba bien escrito como pieza de persuasión. No resistía una revisión de capacidad. Esta tabla es el mapa de cambios. Nada de lo que sigue es un ataque al equipo: es dejar el expediente defendible.",
        align="justify",
    )
    add_table(
        doc,
        ["Tema", "Borrador", "Versión corregida"],
        [
            ["OEE", "Se presenta como metodología aplicada", "Se retira. No hay Disponibilidad × Desempeño × Calidad. Solo se restaron ausencias."],
            ["Nómina", "Capacidad con 9; demanda con almacenistas 9 y 10", "Se trabaja con 9 hasta confirmar. La fila de MP (82,50 HH) son 2 FTE de trabajo, no 2 personas demostradas."],
            ["Jornada", "8:15–16:45 = 8,25 h", "El reloj da 8,50 h. Se usa 8,25 h neta como supuesto (15 min no productivos) y se declara."],
            ["Vacaciones / salud / curva", "Se restan cada semana como si fueran eternos", "Foto de estrés actual (escenario B). Para plantilla se anualiza (escenario C)."],
            ["Demanda", "Horas de turno = carga real", "Es ocupación programada, no volumen × tiempo estándar. Se etiqueta como tal."],
            ["Tienda nueva", "16,25 HH fijos + 12,5 % extra cada trimestre", "Doble conteo. Una tienda o es apertura temporal o entra al run-rate de la red. No ambas."],
            ["Refresh", "40 HH/semana dentro de la plantilla fija y Célula 4 permanente", "Proyecto con fecha de fin (~12 semanas a 40 HH si ~3 min/pieza). Temporales, no 2 fijos."],
            ["“3 es el número exacto”", "Salto de 2,5 FTE + crecimiento", "No se sostiene. 3 al 77,7 % = 96 HH < 104 de brecha del propio borrador."],
            ["Jueves y viernes", "Tareas estructurales “completamente desatendidas”", "PT ya no hace picking masivo J–V: ~99 HH ya existen en el turno. Hay que ver en qué se van."],
            ["ROI", "“Se autofinancia con 10.000 piezas”", "No hay valor, % recuperable ni costo fully loaded. Se deja la fórmula, no el eslogan."],
            ["OTIF < 75 %", "Riesgo de no actuar", "Sin baseline no es un pronóstico. Se mide 4 semanas y después se fija meta."],
            ["Custodia legal + self-audit", "El equipo es “guardián legal” y se cuenta a sí mismo", "Responsabilidad operativa sí; fianza laboral no. Conteo independiente obligatorio."],
        ],
        [3.3, 6.6, 6.6],
    )

    # ===== 3. HECHOS / SUPUESTOS =====
    add_heading_custom(doc, "3. Hechos, supuestos y datos que faltan", 1)
    add_para(
        doc,
        "Toda cifra de este documento entra en una de tres cajas. Lo que está en ámbar no puede usarse como si fuera un dato de WMS.",
        align="justify",
    )
    add_table(
        doc,
        ["Hecho (sale del borrador o del reloj)", "Supuesto explícito (se puede cambiar)", "Falta — no decidir plantilla fija sin esto"],
        [
            ["Turno declarado 8:15 a 16:45", "Jornada neta 8,25 h (15 min no productivos)", "Nómina real: ¿9 o 10 personas hoy, con nombres y fechas?"],
            ["Carga asignada de 287,75 HH/semana", "Semana laboral de 5 días", "Volúmenes 12 semanas: pedidos web, líneas tienda, bultos, recepciones MP"],
            ["Picking masivo de PT en L–M–Mi", "Conteos 30 HH, 5S 15 HH, devoluciones 20 HH", "Qué hacen realmente las células PT el jueves y el viernes"],
            ["Plan verbal de 1 tienda / 3 meses y 8 tiendas en PT", "Vacaciones 15 días hábiles/año y ausentismo 4 %", "Plan de aperturas firmado (fechas) y restricción de camión/muelle"],
            ["Proyecto Refresh de 10.000 piezas", "3 minutos por pieza para dimensionar el proyecto", "Valor, aging y % recuperable de esas 10.000 piezas"],
            ["Hay un analista de inventario (Jeremy Urenë) citado para discrepancias", "E-commerce no escala lineal con tiendas (se deja plano)", "OTIF, ERI, error de picking y horas extra de los últimos 12 meses"],
        ],
        [5.5, 5.5, 5.5],
    )

    # ===== 4. CAPACIDAD =====
    add_heading_custom(doc, "4. Capacidad: tres escenarios, no uno", 1)
    add_heading_custom(doc, "4.1 Jornada", 2)
    add_para(
        doc,
        "De 08:15 a 16:45 hay 8,50 horas de reloj. El borrador usó 8,25 horas. La diferencia es 0,25 h × 9 personas × 5 días = 11,25 HH/semana (casi 0,3 FTE) antes de cualquier ajuste. Hasta que RR.HH. confirme si hay descanso no remunerado, este documento usa 8,25 h como jornada neta informada.",
        align="justify",
    )
    add_table(
        doc,
        ["Concepto", "Cálculo", "Resultado"],
        [
            ["Horas netas / persona / día", "Dato informado", "8,25 h"],
            ["Horas netas / persona / semana", "8,25 × 5", "41,25 HH"],
            ["Capacidad teórica de 9 almacenistas", "9 × 41,25", "371,25 HH/semana"],
            ["Si la jornada neta fuera 8,50 h", "9 × 8,50 × 5", "382,50 HH/semana"],
        ],
        [6.5, 5.0, 5.0],
    )

    add_heading_custom(doc, "4.2 Escenario A — Teórica bruta", 2)
    add_para(doc, "371,25 HH/semana. Es el techo de reloj, no lo que se puede planificar. Sirve de ancla, no de meta de utilización al 100 %.", align="justify")

    add_heading_custom(doc, "4.3 Escenario B — Foto de esta semana (el del borrador)", 2)
    add_para(
        doc,
        "Sirve para explicar por qué el equipo siente saturación hoy. No sirve para crear tres cargos indefinidos. Las tres restas son eventos puntuales o de un subconjunto de personas:",
        align="justify",
    )
    add_table(
        doc,
        ["Resta", "Cálculo", "HH / semana", "Por qué no es estructural"],
        [
            ["1 persona de vacaciones (PT Manufacturado)", "1 × 41,25", "−41,25", "Las vacaciones se anualizan en todo el equipo. 1 FTE ausente todas las semanas equivaldría a ~4 semanas × 9 personas / año, muy por encima de un calendario normal."],
            ["2 personas al 30 % por salud (Equipamiento)", "2 × 41,25 × 30 %", "−24,75", "30 % crónico es alerta de SST / RR.HH., no un parámetro de plantilla. Contratar no cura la causa."],
            ["1 persona nueva en MP al 40 % de curva", "1 × 41,25 × 40 %", "−16,50", "Es temporal. Tres contrataciones nuevas empeoran la curva el primer trimestre porque el experto deja de operar para formar."],
            [("Total restas / capacidad “real” del borrador", "371,25 − 82,50", "288,75  (77,8 %)", "Coincidencia incómoda: las restas (82,50) casi igualan la holgura teórica (83,50) y dejan el día a día en 99,6 %."), AMBER_BG, True],
        ],
        [4.4, 3.3, 2.6, 6.2],
    )

    add_heading_custom(doc, "4.4 Escenario C — Estructural (el que se usa para decidir plantilla)", 2)
    add_para(
        doc,
        "Factores anuales, aplicables a las 9 personas, etiquetados como supuesto. Si RR.HH. entrega 20 días de vacaciones o un ausentismo real distinto, se recalcula en 10 minutos. No se mete la curva de aprendizaje ni las dos bajas al 30 %.",
        align="justify",
    )
    add_table(
        doc,
        ["Factor", "Premisa", "Cálculo", "Efecto"],
        [
            ["Vacaciones", "15 días hábiles / persona / año  (15/260 = 5,77 %)", "371,25 × (1 − 15/260) = 349,83", "−21,42 HH/semana"],
            ["Ausentismo recurrente", "4 % sobre la capacidad post-vacaciones (a validar con RR.HH.)", "349,83 × 0,96 = 335,84", "−13,99 HH/semana"],
            ["Curva de aprendizaje", "0 % en el techo estructural", "No aplica", "Se trata en el plan de onboarding, no en el techo"],
            [("Capacidad estructural C", "9 personas, semana de 5 días", "371,25 × 0,9423 × 0,96", "335,84 HH  =  8,14 FTE"), GREEN_BG, True],
        ],
        [3.4, 5.2, 4.6, 3.3],
    )
    add_para(
        doc,
        "Sensibilidad (para no pelearse por un decimal): si las vacaciones fueran 20 días y el ausentismo 6 %, la capacidad C baja a 371,25 × (240/260) × 0,94 ≈ 322,3 HH. La brecha run-rate sube de 0,41 a ~0,74 FTE. Sigue sin ser 2,5 FTE.",
        size=10.5,
        italic=True,
        color=MUTED,
    )

    # ===== 5. DEMANDA =====
    add_heading_custom(doc, "5. Demanda de trabajo", 1)
    add_heading_custom(doc, "5.1 Advertencia de método", 2)
    add_callout(
        doc,
        "Esto no es un estudio de tiempos",
        "La tabla de 287,75 HH es el horario actual (“L–V de 9:00 a 12:00”), no “X pedidos × Y minutos estándar”. Si el picking es lento, el slotting es malo o no hay WMS, contratar 3 personas convierte esa ineficiencia en costo fijo. La petición de este documento se sostiene con esa reserva: primero proteger horas de control y el proyecto Refresh; en paralelo, medir volumen y productividad 4 semanas.",
        AMBER_BG,
    )

    add_heading_custom(doc, "5.2 Carga del día a día (ejecución actual)", 2)
    add_table(
        doc,
        ["Frente", "Asignación en el borrador", "Detalle", "HH / semana", "Nota de revisión"],
        [
            ["Pedidos web + cuadro chat", "Almacenista 1", "L–V 9:00–12:00 (3 h/día)", "15,00", "Solo mañana. El SLA < 24 h de tarde/viernes no está cubierto con horas; es un problema de cutoff, no solo de FTE."],
            ["Reabastecimiento tienda nueva", "Almacenista 1", "L–V 13:00–16:15 (3,25 h/día)", "16,25", "No puede ser eterno y además un +12,5 % trimestral. Se reclasifica: apertura o 9.ª tienda, no las dos."],
            ["Consumibles (tiendas + taller/backoffice)", "Almacenista 2", "L–M jornada completa + Mi–V 9–12", "25,50", "Aritmética correcta (16,50 + 9,00)."],
            ["PT Manufacturado (8 tiendas)", "Almacenistas 3, 4 y 5", "L–M–Mi jornadas completas (3 × 24,75)", "74,25", "Jueves y viernes ya quedan libres de picking masivo."],
            ["PT Equipamiento (8 tiendas)", "Almacenistas 6, 7 y 8", "L–M–Mi jornadas completas (3 × 24,75)", "74,25", "Aquí estaban las 2 bajas al 30 %. Reducir esta célula a 2 fijos, como proponía el borrador, choca con su propio diagnóstico."],
            ["Materias primas e insumos", "Almacenistas 9 y 10", "Despacho diario + recepción, conteo y calidad", "82,50", "Inconsistencia 9 vs 10. 82,50 HH = 2 FTE de trabajo. Si solo hay 9 en nómina, esta fila o está sobreestimada o ya hay sobrecarga real en MP."],
            [("SUBTOTAL DÍA A DÍA", "", "", "287,75", "Suma verificada: 15+16,25+25,5+74,25+74,25+82,5."), GREEN_BG, True],
        ],
        [3.4, 2.8, 3.4, 1.8, 5.1],
    )

    add_heading_custom(doc, "5.3 Tareas hoy omitidas: run-rate versus proyecto", 2)
    add_para(
        doc,
        "El borrador sumó 105 HH y las trató como déficit permanente. Hay que partirlas. Las horas de conteos, 5S y devoluciones no tienen memoria de cálculo (no hay SKU, frecuencia ABC ni volumen de retornos). Se conservan como supuesto de planificación y se marcan para validar con una muestra de 2 semanas.",
        align="justify",
    )
    add_table(
        doc,
        ["Tarea", "HH / semana", "Tipo", "Cómo se usa en esta propuesta"],
        [
            ["Conteos cíclicos ABC (MP, PT, equipamiento, consumibles)", "30,00", "Run-rate", "Entra a la demanda sostenible. Cada viernes no puede ser solo self-audit del custodio."],
            ["5S y reordenamiento (mantenimiento, no el kick-off)", "15,00", "Run-rate", "Si 5S aún no está implantado, una parte es proyecto. Aquí se deja como mantenimiento semanal."],
            ["Control de devoluciones (logística inversa)", "20,00", "Run-rate", "Debería crecer con la red de tiendas. Hoy no hay volumen."],
            [("Subtotal sostenible omitido", "65,00", "Run-rate", "287,75 + 65,00 = 352,75 HH de demanda estructural."), GREEN_BG, True],
            [("Proyecto Refresh (10.000 piezas)", "40,00", "Proyecto", "No se capitaliza como cargo fijo. Ver 5.4."), AMBER_BG, True],
        ],
        [5.8, 2.2, 2.3, 6.2],
    )

    add_heading_custom(doc, "5.4 Refresh: dimensionar el proyecto, no la nómina", 2)
    add_para(
        doc,
        "Sin tiempo por pieza el 40 HH/semana es un número redondo. Para no dejarlo en el aire, se usa un rango explícito. Dirección puede sustituir 3 minutos por el dato real de una muestra de 50 piezas.",
        align="justify",
    )
    add_table(
        doc,
        ["Minutos / pieza", "Horas totales (10.000 pzas)", "Semanas a 40 HH/sem", "Semanas con 2 personas dedicadas (82,5 HH)"],
        [
            ["2 min", "333 HH", "8,3 semanas", "4,0 semanas"],
            [("3 min (caso base)", "500 HH", "12,5 semanas", "6,1 semanas"), BLUE_BG, True],
            ["5 min", "833 HH", "20,8 semanas", "10,1 semanas"],
            ["8 min (si hay inspección y relabel pesado)", "1.333 HH", "33,3 semanas", "16,2 semanas"],
        ],
        [4.2, 4.1, 4.1, 4.1],
    )
    add_para(
        doc,
        "Incluso en el caso pesimista de 8 minutos, Refresh es un proyecto de un trimestre con 2 personas dedicadas, no una célula eterna. Al terminar, esas horas desaparecen. Por eso la Célula 4 del borrador (2 fijos) quedaría, post-proyecto, con ~20 HH de devoluciones: una persona de más.",
        align="justify",
    )

    add_heading_custom(doc, "5.5 Demanda que sí se usa para decidir", 2)
    add_table(
        doc,
        ["Escenario de demanda", "Composición", "HH / semana", "FTE a 41,25 h"],
        [
            ["D1 — Solo día a día", "287,75", "287,75", "6,98"],
            [("D2 — Sostenible (recomendado para plantilla)", "287,75 + 65,00", "352,75", "8,55"), GREEN_BG, True],
            ["D3 — Sostenible + Refresh (solo mientras dure el proyecto)", "352,75 + 40,00", "392,75", "9,52"],
        ],
        [6.0, 4.3, 3.1, 3.1],
    )

    # ===== 6. BRECHA =====
    add_heading_custom(doc, "6. Brecha de capacidad (el argumento técnico, corregido)", 1)
    add_para(
        doc,
        "Cifra positiva = déficit (faltan horas). Cifra negativa = holgura. FTE = HH / 41,25. La celda verde es la que se recomienda usar para cargos permanentes. La celda del borrador (D3 contra foto B) es la que producía “2,5 operadores”.",
        align="justify",
    )
    add_table(
        doc,
        ["Demanda ↓  /  Capacidad →", "A teórica  371,25", "C estructural  335,84", "B foto semanal  288,75"],
        [
            ["D1 día a día  287,75", "−83,50 HH  (−2,02 FTE)", "−48,09 HH  (−1,17 FTE)", "−1,00 HH  (~0 FTE)"],
            [("D2 sostenible  352,75", "−18,50 HH  (−0,45 FTE)", "+16,91 HH  (+0,41 FTE)", "+64,00 HH  (+1,55 FTE)"), GREEN_BG, True],
            ["D3 + Refresh  392,75", "+21,50 HH  (+0,52 FTE)", "+56,91 HH  (+1,38 FTE)", ("+104,00 HH  (+2,52 FTE)", AMBER_BG, True)],
        ],
        [4.5, 4.0, 4.0, 4.0],
    )
    add_bullet(doc, " El día a día cabe en 9 personas incluso con la foto mala de esta semana. El problema no es “no alcanzamos a despachar tiendas”; es que el control y el proyecto se caen.")
    add_bullet(doc, " Para plantilla permanente, la brecha honesta es ~0,4 FTE (D2 vs C). Un permanente cubre esa brecha, la curva de MP y un colchón de onboarding.")
    add_bullet(doc, " Refresh explica ~1 FTE adicional por un trimestre. Eso son temporales o un sprint con jueves/viernes de PT, no 2 indefinidos.")
    add_bullet(doc, " El 2,5 FTE del borrador solo aparece si se congelan vacaciones + salud al 30 % + curva como techo para siempre y se trata Refresh como eterno. Ese cruce no se recomienda.")

    add_heading_custom(doc, "6.1 Capacidad oculta de jueves y viernes", 2)
    add_para(
        doc,
        "Seis personas de PT Manufacturado y Equipamiento ya tienen L–M–Mi de picking masivo. Jueves y viernes, si el turno existe, hay 6 × 2 × 8,25 = 99 HH semanales. Las 105 HH “omitidas” del borrador casi caben ahí. Antes de hablar de colapso de OTIF hay que responder, con evidencia de dos semanas, qué ocurre esos dos días: ¿camión, incendio operativo, conteos informales, o tiempo no protegido que se evapora?",
        align="justify",
    )
    add_callout(
        doc,
        "Restricción que el headcount no resuelve",
        "Si el surtido a tienda sale solo L–M–Mi porque así está la ruta del camión, agregar gente el jueves no mejora el OTIF de tiendas. El cuello sería muelle, transporte o cutoff de pedido, no FTE. Eso se confirma con Tráfico / Distribución en la quincena de diagnóstico.",
        BLUE_BG,
    )

    # ===== 7. CRECIMIENTO =====
    add_heading_custom(doc, "7. Crecimiento de tiendas, sin doble conteo", 1)
    add_para(
        doc,
        "Regla: una tienda nueva deja de ser “proyecto de apertura” el día en que entra al reabastecimiento regular. A partir de ahí vive en el run-rate de la red. No se suma 16,25 HH eternos y además un 12,5 % cada trimestre sobre la misma base.",
        align="justify",
    )
    add_table(
        doc,
        ["Componente", "Base 8 tiendas (HH/sem)", "Por tienda adicional", "Supuesto"],
        [
            ["PT Manufacturado + Equipamiento", "148,50", "18,56", "Lineal con n.º de tiendas. Optimista si hay economías; pesimista si el camión satura."],
            ["Consumibles", "25,50", "3,19", "Lineal. A validar."],
            ["Devoluciones (las 20 HH omitidas)", "20,00", "2,50", "Lineal. El web se deja plano (15 HH) por falta de forecast."],
            [("Total run-rate sensible a tiendas", "194,00", "24,25  (0,59 FTE)", "Web, MP e I+D no se proyectan: no hay dato."), BLUE_BG, True],
        ],
        [5.0, 3.8, 3.8, 3.9],
    )
    add_para(doc, "Si el plan de 1 tienda cada 3 meses está aprobado y se parte de 8 tiendas en PT:", size=11, bold=True, color=NAVY, space_after=6)
    add_table(
        doc,
        ["Momento", "Tiendas en red PT", "HH extra vs. hoy (8 tiendas)", "FTE extra", "Implicación de plantilla"],
        [
            ["Hoy", "8", "0,00", "0,00", "No contratar crecimiento por adelantado"],
            ["Cierre T1", "9", "24,25", "0,59", "Caben en el permanente + mejor uso J–V, si el camión aguanta"],
            ["Cierre T2 (tienda 10)", "10", "48,50", "1,18", "Gatillo: evaluar 2.ª vacante permanente en Célula 2"],
            ["Cierre T3", "11", "72,75", "1,76", "Solo si las aperturas T1–T2 ocurrieron de verdad"],
            ["Cierre T4 (año)", "12", "97,00", "2,35", "Pico de año. No se contrata hoy el FTE del mes 12"],
        ],
        [3.0, 3.0, 3.5, 2.2, 4.8],
    )
    add_para(
        doc,
        "Leer el año completo como “faltan 2,35 FTE desde el día 1” es el mismo error del borrador, con otro disfraz. Se contrata contra tiendas abiertas, no contra un roadmap. La línea de 16,25 HH de “tienda nueva” del borrador, si ya es la 9.ª tienda en operación, ya está dentro del paso T1 y no se vuelve a sumar.",
        align="justify",
    )

    # ===== 8. OPCIONES =====
    add_heading_custom(doc, "8. Opciones para Dirección", 1)
    add_para(
        doc,
        "Cuatro caminos. El costo se deja en fórmula porque este expediente no tiene el salario fully loaded ni las extra actuales. Operaciones + RR.HH. pegan el número en la celda y el ranking no cambia.",
        align="justify",
    )
    add_table(
        doc,
        ["Opción", "Qué se hace", "Qué cubre", "Qué no cubre", "Cuándo elegirla"],
        [
            ["0. Solo reorganizar", "4 células, vales, accesos, proteger J–V para conteos, 5S y un avance de Refresh con gente actual", "Ownership, control básico, visibilidad", "Pico Refresh y la brecha de 0,4 FTE si J–V ya están ocupados de verdad", "Si en 15 días se ve que J–V están libres y Refresh puede esperar"],
            [("1. Reorganizar + 2 temporales 4–6 meses", "Opción 0 + sprint Refresh con fecha de fin", "Las 40 HH de proyecto y un colchón a las bajas puntuales", "La brecha estructural de 0,4 FTE y MP/curva", "Si Refresh es la urgencia de caja y MP está estable"), GREEN_BG],
            [("2. Recomendada: 1 permanente + 2 temporales", "Opción 1 + 1 Operador Integral (prioridad Célula 1 MP)", "Brecha 0,4 FTE, MP, onboarding, Refresh con fin, gobernanza", "El pico de la tienda 12; no debe cubrirlo", "Caso base. Pedido de este documento."), GREEN_BG, True],
            ["3. 3 permanentes (borrador)", "12 almacenistas fijos desde el día 1, Célula 4 eterna", "Holgura amplia si se cree el escenario B y las 4 aperturas", "Riesgo de subutilización post-Refresh y costo fijo de más", "Solo si Dirección asume por escrito que vacaciones+salud 30 %+curva son el techo eterno y hay 4 aperturas firmadas. No se recomienda."],
        ],
        [2.6, 3.8, 3.4, 3.4, 3.3],
    )
    add_heading_custom(doc, "Cómo se autofinancia (marco, no eslogan)", 2)
    add_para(
        doc,
        "Costo anual opción 2 ≈ (1 × costo fully loaded permanente) + (2 × costo temporal × meses/12). Beneficio ≈ (valor de piezas Refresh × % recuperable realizado en el plazo) + extra evitada + merma evitada por conteos. Si el valor recuperable de Refresh no supera el costo de los 2 temporales en el plazo del proyecto, Refresh se replantea (liquidar, donar, destruir con acta) en lugar de inflar nómina. Eso es un ROI. Lo anterior no lo era.",
        align="justify",
    )

    # ===== 9. CELULAS =====
    add_heading_custom(doc, "9. Modelo operativo: 4 células, por fases", 1)
    add_para(
        doc,
        "La idea de células especializadas del borrador se mantiene: reduce la responsabilidad diluida y acelera OTIF si el layout acompaña. Lo que se corrige es esperar a 12 personas para empezar. El mapa se implanta con 9.",
        align="justify",
    )
    add_heading_custom(doc, "Fase 0 — Inmediata, 9 personas (esta semana)", 2)
    add_table(
        doc,
        ["Célula", "FTE", "Alcance", "Por qué así con 9"],
        [
            ["1. Materias primas e insumos", "2", "Telas, hilos, insumos, jaula de consumibles (custodia física)", "El experto forma y audita. No se baja de 2: la fila de demanda es 82,50 HH."],
            ["2. PT manufacturado, equipamiento y rampa", "4", "2 en manufacturado + 2 en equipamiento volumétrico + bahía/rampa", "Se baja de 6 nombres del roster a 4 fijos de célula; L–M–Mi picking, J–V conteo/5S/apoyo Refresh."],
            ["3. Fulfillment e-commerce, chat y entrega de consumibles", "2", "Picking web/chat, empaque B2C, entrega de consumibles con vale a Célula 1", "15 + 25,50 = 40,50 HH de run-rate. 2 personas alcanzan si hay cutoff de web."],
            ["4. Logística inversa", "1", "Devoluciones y garantías; Refresh solo en la medida que sobre tiempo", "20 HH de devoluciones no justifican 2 fijos. Refresh espera a los temporales o a J–V de Célula 2."],
        ],
        [4.0, 1.4, 5.6, 5.5],
    )
    add_heading_custom(doc, "Fase 1 — Si se aprueba la petición (1 fijo + 2 temporales)", 2)
    add_table(
        doc,
        ["Célula", "FTE", "Cambio"],
        [
            ["1. MP e insumos", "3", "Entra el permanente. Baja la curva, hay doble control en entrega a Producción y ERI de la jaula de consumibles."],
            ["2. PT + rampa", "4", "Sin cambio de headcount. Se protegen jueves y viernes. No se reduce Equipamiento a 2 mientras haya bajas de salud abiertas."],
            ["3. Web + consumibles (ejecución)", "2", "Cutoff escrito de e-commerce. Entrega de consumibles contra vale; no son dueños del stock de la jaula."],
            [("4. Inversa + Refresh", "1 fijo + 2 temporales", "Los 2 temporales tienen meta de piezas/semana y fecha de fin (caso base: 12 semanas a ~3 min/pieza). El fijo se queda con devoluciones al terminar."), GREEN_BG, True],
        ],
        [4.2, 3.3, 9.0],
    )
    add_heading_custom(doc, "Fase 2 — Gatillo de la tienda 10", 2)
    add_para(
        doc,
        "Solo si la 9.ª y la 10.ª tienda abrieron y el OTIF de tiendas o las HH reales de Célula 2 superan el estándar. Entonces se evalúa la 2.ª permanente en Célula 2 (PT), no antes. La Célula 4 no se infla post-Refresh.",
        align="justify",
    )

    add_heading_custom(doc, "Conflicto del borrador que se cierra: consumibles", 2)
    add_para(
        doc,
        "En el texto original la Célula 1 era “custodio de la jaula de consumibles” y la Célula 3 “entrega consumibles”. Eso diluye el ERI. Aquí la regla es única: el stock de la jaula es de Célula 1. Célula 3 ejecuta la entrega a tienda/taller con vale firmado, igual que Producción no entra a MP sin orden. Un dueño de inventario, un ejecutor de movimiento.",
        align="justify",
    )

    # ===== 10. GOBERNANZA =====
    add_heading_custom(doc, "10. Gobernanza de inventario (sin teatro legal)", 1)
    add_para(
        doc,
        "Se conserva lo útil del borrador (vales, perímetro, dueño por nave). Se retira “guardián legal / control absoluto”. En la mayoría de marcos laborales no se puede convertir al almacenista en fiador de la merma. Pedirlo por escrito crea riesgo con RR.HH. y no mejora un solo conteo.",
        align="justify",
    )
    add_table(
        doc,
        ["Regla", "Cómo se opera", "Qué se evita"],
        [
            ["Dueño operativo por nave", "Cada célula firma un acta simple de responsabilidad de ubicación (orden, ERI, acceso). No es fianza.", "Castigo salarial por diferencia de inventario"],
            ["Movimiento solo con vale / documento", "Célula 1 entrega a Producción con OP. Célula 4 transfiere a Célula 2 solo con reclasificación en ERP/WMS y auditoría.", "Salidas “de confianza” y Refresh que se evapora"],
            ["Acceso restringido", "Llave o tarjeta por nave (MP, consumibles, PT, Refresh). Visita de Producción / I+D / comercial con registro.", "Tránsito libre que hace inútil cualquier KPI de ERI"],
            ["Conteo dual", "Viernes: 90 minutos de recuento dirigido (la célula cuenta un cuadrante que NO eligió ella). Independiente: el analista de inventario (Jeremy Urenë o quien designe Control) sortea SKU A y cuenta a ciegas.", "Self-audit como único control (el custodio se audita a sí mismo)"],
            ["Refresh", "Prohibida la salida sin acta de reclasificación. Meta semanal pública de piezas procesadas / valorizadas / dadas de baja.", "Proyecto eterno sin throughput"],
        ],
        [3.6, 7.5, 5.4],
    )

    add_heading_custom(doc, "KPIs: primero baseline, después meta", 2)
    add_para(
        doc,
        "El borrador ponía OTIF > 95 %, ERI > 98 %, picking < 0,5 % y web < 24 h sin decir de dónde partimos. Un “si no contratamos, OTIF cae bajo 75 %” sin serie histórica es presión, no pronóstico. El régimen de este documento es: 4 semanas de medición, luego meta.",
        align="justify",
    )
    add_table(
        doc,
        ["KPI", "Quién lo mueve", "Cómo se mide (semana 1–4)", "Meta (a fijar en día 30, no antes)"],
        [
            ["Utilización HH (reloj vs. asignado vs. estándar)", "Jefatura almacén", "Planilla de horas por célula, diaria", "Techo de planificación, no 100 %"],
            ["OTIF tiendas", "Célula 2 + transporte", "Pedidos completos a tiempo / pedidos", "Se fija contra el baseline real"],
            ["Despacho web y error de picking", "Célula 3", "% salidas en cutoff; % reclamos / líneas", "Cutoff escrito; error a la baja vs. mes 1"],
            ["ERI por nave", "Cada célula + conteo independiente", "Exactitud de registro en muestra ABC", "Mejora vs. primera foto, no un 98 % inventado"],
            ["Throughput Refresh", "Célula 4 + temporales", "Piezas/semana y valor realizado", "Curva que cierre el proyecto en el plazo firmado"],
            ["Ausentismo y extra", "RR.HH. + jefatura", "%, causas, HH extra", "Bajar extra de incendio, no maquillar salud al 30 %"],
        ],
        [4.0, 3.2, 4.6, 4.7],
    )

    # ===== 11. PLAN 90 DIAS =====
    add_heading_custom(doc, "11. Plan de 90 días", 1)
    add_table(
        doc,
        ["Cuándo", "Qué", "Salida visible"],
        [
            ["Días 1–15", "Cerrar nómina 9 vs 10. Estudio de tiempos jueves/viernes (qué se hizo, cada hora, 10 días hábiles). Extraer OTIF, ERI, extra, ausentismo 12 meses. Muestra de 50 piezas Refresh (minutos y valor). Confirmar con Comercial las fechas de apertura.", "Carpeta de evidencia. Recalcular escenarios C y D2 si un supuesto cambia >10 %."],
            ["Días 1–21", "Levantar las 4 células con 9 personas. Vales, llaves, acta de dueño de nave. Cutoff e-commerce por escrito. Viernes de conteo dual.", "Mapa físico en la pared y en el ERP (si existe). Primera foto de ERI."],
            ["Día 21", "Decisión formal de la petición (1 fijo + 2 temporales) con los datos de la quincena, no contra el borrador.", "Acta de Dirección. Si J–V estaban libres y Refresh vale poco, se puede bajar a opción 1 o 0."],
            ["Días 22–90", "Onboarding del permanente en Célula 1. Sprint Refresh con temporales y meta semanal. Proteger J–V de Célula 2 (conteos/5S). No usar esos días para “lo que vaya saliendo”.", "Curva de piezas Refresh. Segundo ERI. Informe de extra."],
            ["Día 90", "Go / no-go de la 2.ª permanente: solo si hubo apertura real hacia la tienda 10 o si Célula 2 está sobre estándar con evidencia.", "Una página a Dirección. Los temporales de Refresh se cierran o se prorrogan con nueva fecha, nunca se silencian a fijos."],
        ],
        [2.6, 6.5, 7.4],
    )

    # ===== 12. RIESGOS =====
    add_heading_custom(doc, "12. Riesgos de no actuar y de actuar mal", 1)
    add_table(
        doc,
        ["Riesgo", "Si no se hace nada", "Si se hace lo del borrador (3 fijos ya)"],
        [
            ["Control de inventario", "Conteos y Refresh siguen sin dueño. El muerto de 10.000 piezas no se mueve.", "Hay gente, pero el self-audit y la “custodia legal” no cierran merma y crean ruido laboral."],
            ["Servicio a tiendas / web", "Si J–V ya están quemados en incendios, el crecimiento de tiendas sí aprieta OTIF. Eso hay que medirlo, no afirmar <75 %.", "Se paga FTE de más y el OTIF no sube si el cuello es el camión o el cutoff."],
            ["Costo", "Extra y desorden. Costo oculto.", "Costo fijo de 3 personas + curva de 3 novatos a la vez, con Refresh terminado a mitad de año y Célula 4 sobrada."],
            ["MP y Producción", "La curva del elemento nuevo y las entregas sin vale siguen.", "Si el permanente no entra a MP, el problema original queda igual mientras se infla PT o Refresh."],
        ],
        [3.3, 6.6, 6.6],
    )

    # ===== 13. DECISION =====
    add_heading_custom(doc, "13. Hoja de decisión (para firmar)", 1)
    add_para(
        doc,
        "Marcar lo que se aprueba. Lo no marcado no se ejecuta. Las filas verdes (puntos 3 y 4) son el paquete de headcount que este expediente considera coherente.",
        align="justify",
    )
    add_table(
        doc,
        ["N.º", "Decisión", "Sí / No / Fecha"],
        [
            ["1", "Se aprueba implantar 4 células con las 9 personas actuales (Fase 0), vales y acceso restringido.", ""],
            ["2", "Se retira del discurso interno el OEE no calculado, el OTIF <75 % sin baseline y la custodia legal / self-audit como único control.", ""],
            [("3", "Se aprueba 1 vacante permanente de Operador Integral de Almacén, incorporación a Célula 1 (MP).", ""), GREEN_BG, True],
            [("4", "Se aprueban 2 cupos temporales (4–6 meses) para Refresh, con meta de piezas/semana y fecha de fin en el contrato o el acta.", ""), GREEN_BG, True],
            ["5", "El plan de 1 tienda / 3 meses queda confirmado con fechas (o se declara que aún no está firmado). La 2.ª permanente no se abre hasta la tienda 10 real.", ""],
            ["6", "En 15 días Operaciones entrega: nómina 9 vs 10, diario J–V, OTIF/ERI/extra 12 meses, muestra Refresh (minutos y valor).", ""],
            ["7", "RR.HH. confirma jornada neta (8,25 vs 8,50), días de vacaciones y ausentismo 12 meses, y el costo fully loaded para cerrar el ROI.", ""],
        ],
        [1.2, 12.3, 3.0],
    )
    add_para(doc, "Firmas", size=12, bold=True, color=NAVY, space_after=8)
    firmas = doc.add_table(rows=3, cols=3)
    firmas.alignment = WD_TABLE_ALIGNMENT.CENTER
    heads = ["Operaciones / Almacén", "RR.HH.", "Dirección"]
    for i, h in enumerate(heads):
        write_cell(firmas.rows[0].cells[i], h, bold=True, size=10, color=WHITE, align="center", fill=HEADER_BG)
        write_cell(firmas.rows[1].cells[i], "\n\nNombre:\nCargo:\nFecha:\nFirma:\n", size=10, color=DARK, fill="FFFFFF")
        write_cell(firmas.rows[2].cells[i], "Aprobado  /  Aprobado con cambios  /  Rechazado", size=8.5, italic=True, color=MUTED, align="center", fill=ROW_ALT)
    set_table_widths(firmas, [5.5, 5.5, 5.5])
    doc.add_paragraph()

    # ===== ANEXO A =====
    add_heading_custom(doc, "Anexo A. Memoria de cálculo (para que cualquiera reprocese)", 1)
    add_para(doc, "Constantes", size=12, bold=True, color=ACCENT, space_after=6)
    add_para(doc, "H_dia = 8,25 h neta informada.  Dias = 5.  N = 9 personas.  H_sem_per = 8,25 × 5 = 41,25.  Teórica A = 9 × 41,25 = 371,25.", size=11)
    add_para(doc, "Foto B (borrador, aritmética correcta, uso prohibido para plantilla fija)", size=12, bold=True, color=ACCENT, space_after=6)
    add_para(
        doc,
        "Vac_foto = 41,25.  Salud_foto = 2 × 41,25 × 0,30 = 24,75.  Curva_foto = 41,25 × 0,40 = 16,50.  Restas = 82,50.  B = 371,25 − 82,50 = 288,75.  288,75 / 371,25 = 77,78 %.",
        size=11,
    )
    add_para(doc, "Estructural C (uso para plantilla)", size=12, bold=True, color=ACCENT, space_after=6)
    add_para(
        doc,
        "Vac_anual = 15/260.  C = 371,25 × (1 − 15/260) × (1 − 0,04) = 371,25 × 0,94230769 × 0,96 = 335,838.  Se publica 335,84 HH.  FTE efectivos = 335,84 / 41,25 = 8,141.",
        size=11,
    )
    add_para(doc, "Demanda", size=12, bold=True, color=ACCENT, space_after=6)
    add_para(
        doc,
        "D1 = 15,00 + 16,25 + 25,50 + 74,25 + 74,25 + 82,50 = 287,75.  Omitidas_run = 30 + 15 + 20 = 65,00.  D2 = 287,75 + 65,00 = 352,75.  Refresh = 40,00.  D3 = 392,75.",
        size=11,
    )
    add_para(doc, "Brechas usadas en el cuerpo (Demanda − Capacidad)", size=12, bold=True, color=ACCENT, space_after=6)
    add_para(
        doc,
        "D2 − C = 352,75 − 335,84 = 16,91 HH = 16,91 / 41,25 = 0,410 FTE.  D3 − B = 392,75 − 288,75 = 104,00 HH = 2,521 FTE (cifra del borrador, no recomendada).  3 contrataciones teóricas = 123,75 HH.  3 × factor B = 123,75 × 288,75 / 371,25 = 96,25 HH < 104.",
        size=11,
    )
    add_para(doc, "Refresh (caso base 3 min)", size=12, bold=True, color=ACCENT, space_after=6)
    add_para(
        doc,
        "10.000 × 3 / 60 = 500 HH.  500 / 40 = 12,5 semanas.  500 / 82,5 = 6,06 semanas con 2 personas a jornada completa.",
        size=11,
    )
    add_para(doc, "Crecimiento por tienda", size=12, bold=True, color=ACCENT, space_after=6)
    add_para(
        doc,
        "PT unitario = 148,50 / 8 = 18,5625.  Consumibles unitario = 25,50 / 8 = 3,1875.  Devoluciones unitario = 20,00 / 8 = 2,50.  Extra por tienda = 24,25 HH = 0,588 FTE.  Extra a 12 tiendas = 4 × 24,25 = 97,00 HH = 2,352 FTE.",
        size=11,
    )
    add_para(doc, "Jueves–viernes PT", size=12, bold=True, color=ACCENT, space_after=6)
    add_para(doc, "6 personas × 2 días × 8,25 = 99,00 HH/semana ya pagadas dentro de la nómina actual de las células de PT.", size=11)

    # ===== ANEXO B =====
    add_heading_custom(doc, "Anexo B. Preguntas que Operaciones debe responder en 15 días", 1)
    add_bullet(doc, " ¿Cuántas personas hay hoy en planilla de almacén, con nombre, fecha de ingreso y si alguna está en licencia? El texto no puede seguir diciendo 9 y listando hasta el 10.")
    add_bullet(doc, " ¿La jornada de 8:15–16:45 incluye almuerzo o descanso? ¿8,25 h es neta pactada?")
    add_bullet(doc, " Entregar el calendario de vacaciones 12 meses y el % de ausentismo real. El 30 % de dos personas no se vuelve parámetro.")
    add_bullet(doc, " Diario de jueves y viernes de las últimas 2 semanas (aunque sea a mano): hora, tarea, persona. Sin eso, las 105 HH omitidas son una lista de deseos.")
    add_bullet(doc, " Volumen 12 semanas: pedidos y líneas web; bultos o líneas a tienda; recepciones MP; devoluciones; consumibles.")
    add_bullet(doc, " OTIF tiendas, fill rate, ERI, error de picking, extra, quiebres. Aunque sea imperfecto. Cero es peor que un número feo.")
    add_bullet(doc, " Las 10.000 piezas: valor a costo, aging, % estimado vendible / donable / merma. Muestra de 50 piezas cronometrada.")
    add_bullet(doc, " ¿El picking L–M–Mi lo manda el camión? Capacidad de muelle y ruta.")
    add_bullet(doc, " Forecast de e-commerce (¿crece independiente de tiendas?).")
    add_bullet(doc, " Plan de aperturas: documento de Comercial/Dirección, no un ritmo deseado.")

    # ===== ANEXO C =====
    add_heading_custom(doc, "Anexo C. Cómo usar este documento en una reunión", 1)
    add_bullet(doc, " No abrir con OEE ni con “irrebatible / clase mundial”. Abrir con la hoja de decisión de la sección 13 y la tabla de brechas de la sección 6.")
    add_bullet(doc, " Si alguien vuelve a 2,5 FTE, pedir que señale la celda: es D3 contra B (Refresh eterno + foto de salud/vacaciones). No es D2 contra C.")
    add_bullet(doc, " Si piden 3 fijos “para no quedarnos cortos con las tiendas”, mostrar la tabla trimestral: el FTE de la tienda 12 no se contrata en el mes 1.")
    add_bullet(doc, " Si piden cero contrataciones, acordar por escrito el diario J–V y una fecha de Refresh. Cero sin medición es dejar el muerto quieto.")
    add_bullet(doc, " Los números en este Word se pueden recalcular con el script docs/generar_propuesta_almacen.py o a mano con el Anexo A. Si cambia un supuesto (vacaciones, minutos/pieza), se cambia el supuesto, no se maquilla el gap.")

    # ===== ANEXO D =====
    add_heading_custom(doc, "Anexo D. Glosario breve", 1)
    add_table(
        doc,
        ["Término", "Significado en este documento"],
        [
            ["HH", "Hora-hombre. Una persona, una hora de jornada neta."],
            ["FTE", "Full-time equivalent. Aquí 1 FTE = 41,25 HH/semana."],
            ["Run-rate", "Carga que se repite todas las semanas (no un proyecto)."],
            ["OEE", "Overall Equipment Effectiveness (máquinas). No se usa: no está calculado."],
            ["OLE", "Overall Labor Effectiveness, análogo de mano de obra. Tampoco se calcula: faltan desempeño y calidad medidos."],
            ["OTIF", "On Time In Full: pedido completo y a tiempo."],
            ["ERI", "Exactitud de registro de inventario."],
            ["5S", "Orden y estandarización del puesto (clasificar, ordenar, limpiar, estandarizar, sostener)."],
            ["Refresh", "Proyecto de saneamiento de ~10.000 piezas inmovilizadas. Tiene fecha de fin."],
            ["Cutoff", "Hora límite para que un pedido web salga el mismo día."],
        ],
        [3.5, 13.0],
    )

    add_para(
        doc,
        "Documento elaborado a partir de la revisión crítica del borrador “Propuesta aumento de plantilla / Reestructuración Team Almacén Fábrica”. Las sumas originales se verificaron una a una. Las conclusiones de headcount se reconstruyeron con capacidad estructural y con Refresh tratado como proyecto. Cualquier cifra ámbar debe sustituirse por dato de sistema o de RR.HH. antes de convertirla en cargo indefinido.",
        size=9.5,
        italic=True,
        color=MUTED,
        space_after=0,
    )

    out_dir = Path("/workspace/docs")
    out_dir.mkdir(parents=True, exist_ok=True)
    path = out_dir / "Propuesta_Capacidad_Almacen_Fabrica.docx"
    doc.save(path)

    artifact = Path("/opt/cursor/artifacts") / "Propuesta_Capacidad_Almacen_Fabrica.docx"
    artifact.parent.mkdir(parents=True, exist_ok=True)
    doc.save(artifact)
    return path


if __name__ == "__main__":
    p = build()
    print(p)
    print("bytes", p.stat().st_size)
