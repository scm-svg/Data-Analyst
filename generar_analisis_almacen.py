#!/usr/bin/env python3
"""Genera documento Word con revisión crítica de la propuesta de almacén."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, fill_hex: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc, headers, rows, header_fill="1F4E79"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)
        set_cell_shading(hdr[i], header_fill)

    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()
    return table


def add_bullet_list(doc, items, level=0):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.25 * (level + 1))


def build_document():
    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # --- PORTADA ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "REVISIÓN CRÍTICA Y ANÁLISIS TÉCNICO\n"
        "Propuesta de Aumento de Plantilla\n"
        "Reestructuración Team Almacén Fábrica"
    )
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(31, 78, 121)

    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Documento de análisis ejecutivo\nCon verificación matemática, observaciones y recomendaciones")
    r.font.size = Pt(12)
    r.italic = True

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m = meta.add_run("Fecha de elaboración: 26 de agosto de 2026\nTipo: Revisión independiente de propuesta operativa")
    m.font.size = Pt(11)

    doc.add_page_break()

    # --- 1. RESUMEN EJECUTIVO ---
    doc.add_heading("1. Resumen ejecutivo", level=1)
    doc.add_paragraph(
        "Este documento presenta una revisión crítica e independiente de la propuesta titulada "
        "«Propuesta Aumento de Plantilla – Reestructuración Team Almacén Fábrica», cuyo objetivo "
        "principal es justificar el incremento de la plantilla de 9 a 12 almacenistas mediante "
        "metodología de capacidad instalada vs. demanda y reorganización en células de trabajo."
    )
    doc.add_paragraph(
        "Tras verificar los cálculos, contrastar supuestos y evaluar coherencia del planteamiento, "
        "se concluye que la propuesta tiene una base estratégica sólida, pero requiere correcciones "
        "metodológicas y respaldo de datos antes de presentarse a Gerencia o Junta Directiva."
    )

    doc.add_heading("Veredicto general", level=2)
    add_table(
        doc,
        ["Dimensión", "Calificación", "Comentario"],
        [
            ["Planteamiento del problema", "Bueno", "Capacidad vs. demanda claramente identificada"],
            ["Estructura narrativa", "Muy bueno", "Flujo lógico: diagnóstico → brecha → solución → gobernanza"],
            ["Rigor matemático", "Parcial", "Sumas correctas, pero supuestos y mezcla de conceptos"],
            ["Sustento de datos", "Débil", "Faltan fuentes, históricos y time studies"],
            ["Marco OEE", "Débil", "Referencia teórica sin cálculo real de OEE"],
            ["ROI / argumento financiero", "Insuficiente", "Afirmación sin cifras"],
            ["Propuesta de 12 personas", "Razonable", "Células bien pensadas; justificación de «3» mejorable"],
        ],
    )

    p = doc.add_paragraph()
    r = p.add_run("Conclusión principal: ")
    r.bold = True
    p.add_run(
        "Es una buena base operativa y organizacional, pero NO está lista para Junta Directiva "
        "sin corregir inconsistencias, documentar supuestos y cuantificar el retorno de la inversión."
    )

    doc.add_page_break()

    # --- 2. DOCUMENTO REVISADO ---
    doc.add_heading("2. Documento revisado", level=1)
    doc.add_paragraph(
        "La propuesta original plantea que, con una tasa de expansión de 1 tienda nueva cada 3 meses, "
        "la estructura actual de 9 almacenistas entró en déficit crítico de capacidad. Propone:"
    )
    add_bullet_list(
        doc,
        [
            "Calcular capacidad teórica y real (ajustada por vacaciones, salud y curva de aprendizaje).",
            "Cuantificar demanda operativa diaria y tareas omitidas.",
            "Demostrar brecha de 104 horas-hombre (HH) semanales (~2,5 FTE).",
            "Solicitar 3 almacenistas adicionales (total 12) organizados en 4 células.",
            "Implementar modelo de custodia por ubicación con KPIs y gobernanza.",
        ],
    )

    doc.add_heading("Objetivo de esta revisión", level=2)
    add_bullet_list(
        doc,
        [
            "Verificar la exactitud de los cálculos matemáticos.",
            "Cuestionar supuestos y detectar inconsistencias.",
            "Evaluar solidez del marco teórico (OEE, capacidad vs. demanda).",
            "Identificar vacíos de datos y riesgos de credibilidad.",
            "Proponer mejoras concretas para fortalecer la presentación.",
        ],
    )

    # --- 3. REVISIÓN MATEMÁTICA ---
    doc.add_heading("3. Revisión matemática detallada", level=1)

    doc.add_heading("3.1 Capacidad teórica vs. capacidad real", level=2)
    doc.add_paragraph("Cálculos verificados en la propuesta original:")

    add_table(
        doc,
        ["Concepto", "Cálculo", "Resultado", "Estado"],
        [
            ["Horas teóricas semanales", "9 × 8,25 h × 5 días", "371,25 HH", "✓ Correcto"],
            ["Ajuste vacaciones (1 persona)", "8,25 × 5 días", "-41,25 HH", "✓ Aritmética correcta"],
            ["Ajuste salud (2 pers., -30%)", "2 × 41,25 × 0,30", "-24,75 HH", "✓ Correcto"],
            ["Curva aprendizaje (-40%)", "1 × 41,25 × 0,40", "-16,50 HH", "✓ Correcto"],
            ["Capacidad real efectiva", "371,25 - 82,50", "288,75 HH", "✓ Correcto"],
            ["Utilización real", "288,75 ÷ 371,25", "77,7%", "✓ Correcto"],
        ],
    )

    doc.add_heading("Observación crítica sobre los ajustes", level=3)
    doc.add_paragraph(
        "Aunque la aritmética es correcta, los ajustes se aplican como si fueran permanentes "
        "cada semana del año. En la práctica operativa esto sobrestima las pérdidas:"
    )
    add_bullet_list(
        doc,
        [
            "Vacaciones: no hay 1 persona de vacaciones todas las semanas. Debería amortizarse anualmente "
            "(~25 días/año ÷ 52 semanas ≈ 0,48 persona/semana, no 1,0 fijo).",
            "Ausentismo por salud al 30% para 2 personas requiere respaldo documental "
            "(incapacidades, histórico de asistencia de 6-12 meses).",
            "Curva de aprendizaje del 40% es temporal; no debe estar en el modelo permanente "
            "salvo que exista rotación continua de personal nuevo.",
        ],
    )

    doc.add_paragraph(
        "Impacto estimado: si se corrigen vacaciones y aprendizaje como promedios anuales, "
        "la capacidad real podría ascender a aproximadamente 310-320 HH semanales, "
        "reduciendo la brecha de 104 HH a un rango de 75-85 HH (~1,8-2,0 FTE en lugar de 2,5)."
    )

    doc.add_heading("3.2 Demanda operativa diaria", level=2)
    add_table(
        doc,
        ["Tarea / Frente", "Asignación", "Detalle", "HH/Semana", "Verificación"],
        [
            ["Pedidos Web + Chat", "Almacenista 1", "L-V 9:00-12:00 (3h/día)", "15,00", "✓"],
            ["Reabastecimiento tienda nueva", "Almacenista 1", "L-V 13:00-16:15 (3,25h/día)", "16,25", "✓"],
            ["Consumibles", "Almacenista 2", "L-M completo + Mi-V 9-12", "25,50", "✓"],
            ["Reabastecimiento PT Manufacturado", "Almacenistas 3, 4, 5", "L, M, Mi jornada completa", "74,25", "✓"],
            ["Reabastecimiento PT Equipamiento", "Almacenistas 6, 7, 8", "L, M, Mi jornada completa", "74,25", "✓"],
            ["Materias Primas e Insumos", "Almacenistas 9 y 10", "Despacho + recepción diaria", "82,50", "⚠ Ver nota"],
            ["SUBTOTAL carga diaria", "—", "—", "287,75", "✓"],
        ],
    )

    p = doc.add_paragraph()
    r = p.add_run("ERROR CRÍTICO DE CONSISTENCIA: ")
    r.bold = True
    r.font.color.rgb = RGBColor(192, 0, 0)
    p.add_run(
        "La propuesta parte de una nómina de 9 almacenistas, pero en Materias Primas asigna tareas "
        "a «Almacenistas 9 y 10». Esto implica que: (a) ya operan con 10 personas y el título está "
        "desactualizado, o (b) las 82,5 HH de MP no tienen ejecutor dedicado y el déficit real es mayor. "
        "Este punto debe aclararse antes de cualquier presentación."
    )

    doc.add_paragraph(
        "Porcentaje de saturación: 287,75 ÷ 288,75 = 99,6%. "
        "La conclusión de que «matemáticamente es imposible realizar las tareas estructurales» "
        "es aritméticamente correcta bajo los supuestos planteados."
    )

    doc.add_heading("3.3 Tareas omitidas y brecha total", level=2)
    add_table(
        doc,
        ["Tarea omitida", "HH/Semana", "Observación"],
        [
            ["Conteos cíclicos ABC", "30,0", "Sin detalle de SKUs ni frecuencia ABC"],
            ["Protocolo 5S y reordenamiento", "15,0", "Sin alcance por nave definido"],
            ["Control de devoluciones", "20,0", "Sin volumen en unidades/semana"],
            ["Proyecto REFRESH (10.000 piezas)", "40,0", "Requiere validación de tiempos unitarios"],
            ["TOTAL omitidas", "105,0", "—"],
        ],
    )

    add_table(
        doc,
        ["Concepto", "Cálculo", "Resultado"],
        [
            ["Demanda total real", "287,75 + 105,0", "392,75 HH/semana"],
            ["Capacidad real", "—", "288,75 HH/semana"],
            ["Brecha operativa", "392,75 - 288,75", "104,00 HH/semana"],
            ["Equivalente FTE", "104 ÷ 41,25", "2,52 ≈ 2,5 operadores"],
        ],
    )

    doc.add_heading("Validación Proyecto REFRESH", level=3)
    doc.add_paragraph(
        "Con 40 HH/semana asignadas: 10.000 piezas ÷ 40 HH = 250 piezas/semana = 50 piezas/día. "
        "Si cada pieza requiere 10 minutos (etiquetado + reclasificación + registro en sistema), "
        "se necesitarían 8,3 HH/día solo para REFRESH, lo que supera la capacidad de 1 persona. "
        "Se recomienda un time study para validar el estándar de minutos por pieza."
    )

    doc.add_heading("3.4 Justificación de 3 vs. 2,5 operadores", level=2)
    doc.add_paragraph(
        "La propuesta salta de 2,5 FTE (104 HH) a solicitar 3 personas nuevas citando "
        "«1 tienda cada 3 meses = +12,5% trimestral». Sin embargo:"
    )
    add_bullet_list(
        doc,
        [
            "Si el 12,5% se aplica solo sobre reabastecimiento PT (148,5 HH): +18,6 HH/trimestre.",
            "Si se aplica sobre toda la operación (392,75 HH): +49 HH/trimestre, más defendible.",
            "Falta una tabla de proyección trimestral que muestre cómo evoluciona la brecha.",
        ],
    )

    add_table(
        doc,
        ["Escenario", "Capacidad real (HH)", "Brecha (HH)", "FTE necesarios"],
        [
            ["Documento original", "288,75", "104,00", "~2,5"],
            ["Capacidad corregida (promedio anual)", "~315", "~78", "~1,9"],
            ["Corregida + crecimiento 4 tiendas/año", "~315 base", "~128", "~3,1"],
        ],
    )

    doc.add_page_break()

    # --- 4. MARCO OEE ---
    doc.add_heading("4. Crítica al marco teórico OEE", level=1)
    doc.add_paragraph(
        "La propuesta invoca OEE (Overall Equipment Effectiveness) adaptado a personal, "
        "definiendo Disponibilidad × Desempeño × Calidad. Sin embargo, no se calcula OEE real."
    )
    add_bullet_list(
        doc,
        [
            "Se restan pérdidas de forma aditiva (-41,25 -24,75 -16,50), no multiplicativa como exige OEE.",
            "OEE real sería: Disponibilidad × Desempeño × Calidad (ej. 0,89 × 0,96 × 1,0 ≈ 85%).",
            "El 77,7% del documento no corresponde a un cálculo OEE estándar.",
            "No hay línea base histórica de productividad (líneas/hora, unidades procesadas).",
            "No hay estándares de tiempo por tarea que sustenten el factor «Desempeño».",
        ],
    )

    doc.add_heading("Recomendación", level=2)
    doc.add_paragraph(
        "Opción A: Eliminar la referencia OEE y hablar de «Eficiencia Operativa de Personal» "
        "con metodología transparente de ajustes. "
        "Opción B: Calcular OEE real con datos históricos:"
    )
    add_table(
        doc,
        ["Factor OEE", "Valor sugerido", "Fuente requerida"],
        [
            ["Disponibilidad", "89%", "Asistencia últimos 6 meses"],
            ["Desempeño", "85%", "Std vs. real en picking/despacho"],
            ["Calidad", "98%", "Errores de conteo y picking"],
            ["OEE Personal calculado", "74,2%", "Producto de los tres factores"],
        ],
    )

    # --- 5. INCONSISTENCIAS CÉLULAS ---
    doc.add_heading("5. Inconsistencias en la propuesta de células (12 almacenistas)", level=1)

    doc.add_heading("5.1 Solapamiento de Consumibles", level=2)
    add_table(
        doc,
        ["Estado", "Responsable de Consumibles", "Rol"],
        [
            ["Situación actual", "Almacenista 2", "Entrega a tiendas + taller/backoffice"],
            ["Célula 1 (propuesta)", "3 almacenistas", "Custodia física de jaula de consumibles"],
            ["Célula 3 (propuesta)", "2 almacenistas", "Entrega de consumibles a tiendas"],
        ],
    )
    doc.add_paragraph(
        "No está definido quién prepara, quién custodia y quién distribuye. "
        "Riesgo de doble responsabilidad o huecos operativos. Debe explicitarse el flujo: "
        "Custodia (Célula 1) → Preparación → Entrega (Célula 3) con protocolo de transferencia."
    )

    doc.add_heading("5.2 Jueves y Viernes «libres de picking» (Célula 2)", level=2)
    doc.add_paragraph(
        "5 almacenistas × 2 días × 8,25 h = 82,5 HH disponibles para 5S y conteos. "
        "Pero las tareas omitidas solo presupuestan 45 HH (30 conteos + 15 5S). "
        "Hay 37,5 HH de capacidad adicional no contabilizada, o las HH omitidas están subestimadas."
    )

    doc.add_heading("5.3 Distribución propuesta", level=2)
    add_table(
        doc,
        ["Célula", "Personal", "Foco", "Evaluación"],
        [
            ["Célula 1", "3 ops", "MP, Insumos, Consumibles", "Buena — permite mentoring"],
            ["Célula 2", "5 ops", "PT Manufacturado + Equipamiento", "Buena — absorbe ausentismo"],
            ["Célula 3", "2 ops", "E-commerce + Consumibles", "Revisar solapamiento con Célula 1"],
            ["Célula 4", "2 ops", "Devoluciones + REFRESH", "Buena — dedicación exclusiva"],
            ["TOTAL", "12 ops", "—", "3 + 5 + 2 + 2 = 12 ✓"],
        ],
    )

    doc.add_page_break()

    # --- 6. KPIs ---
    doc.add_heading("6. KPIs propuestos: metas sin línea base", level=1)
    add_table(
        doc,
        ["KPI", "Meta propuesta", "Baseline actual", "Estado"],
        [
            ["ERI MP/Consumibles", ">98%", "No reportado", "❌ Falta dato"],
            ["OTIF Tiendas", ">95%", "No reportado (solo riesgo <75%)", "❌ Falta dato"],
            ["Errores de Picking", "<0,5%", "No reportado", "❌ Falta dato"],
            ["Despacho Web", "<24 horas", "No reportado", "❌ Falta dato"],
            ["Recuperación Stock REFRESH", "Porcentaje", "Sin meta numérica", "❌ Incompleto"],
        ],
    )
    doc.add_paragraph(
        "Sin baseline, las metas no demuestran mejora medible. Gerencia preguntará: "
        "«¿Cuál es el OTIF hoy? ¿Cuántos errores de picking tenemos?». "
        "Se recomienda incluir columna «Situación actual» con datos de los últimos 3-6 meses."
    )

    # --- 7. ROI ---
    doc.add_heading("7. Análisis ROI: la afirmación más débil", level=1)
    doc.add_paragraph(
        "La propuesta afirma que «la incorporación de 3 operadores se autofinancia liberando "
        "las 10.000 piezas del Proyecto Refresh». Sin embargo, falta todo el modelo financiero:"
    )

    add_table(
        doc,
        ["Variable financiera", "Fórmula / dato requerido", "Estado"],
        [
            ["Costo anual 3 operadores", "3 × salario mensual × 12 × (1 + cargas sociales)", "❌ No incluido"],
            ["Valor inventario REFRESH", "10.000 piezas × costo promedio unitario", "❌ No incluido"],
            ["Valor recuperable REFRESH", "% recuperable × valor total", "❌ No incluido"],
            ["Merma evitada por conteos", "Histórico de diferencias de inventario × costo", "❌ No incluido"],
            ["Costo horas extra actuales", "HH extra × tarifa × 52 semanas", "❌ No incluido"],
            ["Costo de NO actuar", "Pérdida ventas por quiebre OTIF", "❌ No incluido"],
        ],
    )

    doc.add_heading("Modelo ROI sugerido (plantilla)", level=2)
    doc.add_paragraph(
        "ROI = (Beneficios anuales - Costo anual personal) ÷ Costo anual personal × 100\n\n"
        "Beneficios anuales = Valor REFRESH recuperado + Merma evitada + Ahorro horas extra + "
        "Ventas protegidas por expansión\n\n"
        "Payback (meses) = Costo anual personal ÷ (Beneficios mensuales)\n\n"
        "Esta plantilla debe completarse con cifras reales antes de la presentación."
    )

    # --- 8. RIESGOS ---
    doc.add_heading("8. Riesgos y afirmaciones sin evidencia", level=1)
    add_table(
        doc,
        ["Afirmación en la propuesta", "Riesgo", "Acción recomendada"],
        [
            ["Colapso OTIF <75%", "Sin OTIF actual documentado", "Incluir dashboard OTIF últimos 6 meses"],
            ["1 tienda cada 3 meses", "¿Está en plan oficial de expansión?", "Anexar cronograma de aperturas"],
            ["10.000 piezas muertas REFRESH", "Sin valor en $ ni antigüedad", "Anexar valorización de inventario"],
            ["Baja responsabilidad reportada", "Lenguaje subjetivo", "Reemplazar con datos de incidencias"],
            ["Intermitencia salud 30%", "Sin respaldo", "Anexar histórico de incapacidades"],
            ["SLA despacho web <24h", "Sin cumplimiento actual", "Medir % pedidos despachados <24h hoy"],
        ],
    )

    doc.add_page_break()

    # --- 9. FORTALEZAS ---
    doc.add_heading("9. Fortalezas de la propuesta (mantener y reforzar)", level=1)
    add_bullet_list(
        doc,
        [
            "Estructura narrativa clara: diagnóstico → brecha cuantificada → solución → gobernanza.",
            "Identificación correcta de tareas omitidas críticas (conteos, 5S, devoluciones, REFRESH).",
            "Modelo de células con custodia por ubicación — alineado con estándares WMS de clase mundial.",
            "Self-audit los viernes (2 horas por célula) — práctica concreta y auditable.",
            "Protocolo de entregas inter-áreas con vales — necesario para custodia efectiva.",
            "La brecha de ~104 HH es un argumento defendible si se corrigen supuestos de capacidad.",
            "El argumento de «responsabilidad diluida» es válido y la solución de custodia es coherente.",
        ],
    )

    # --- 10. RECOMENDACIONES ---
    doc.add_heading("10. Recomendaciones concretas", level=1)

    doc.add_heading("10.1 Correcciones obligatorias antes de presentar", level=2)
    add_bullet_list(
        doc,
        [
            "Aclarar inconsistencia 9 vs. 10 almacenistas en Materias Primas.",
            "Separar capacidad promedio anual vs. semana pico (vacaciones concentradas).",
            "Resolver solapamiento de Consumibles entre Células 1 y 3 con flujo explícito.",
            "Eliminar referencia OEE o calcular OEE real con datos históricos.",
            "Reemplazar afirmaciones subjetivas («baja responsabilidad») por métricas.",
        ],
    )

    doc.add_heading("10.2 Anexos que deben agregarse", level=2)
    add_bullet_list(
        doc,
        [
            "Time study de tareas omitidas (muestreo mínimo de 1 semana).",
            "Histórico de asistencia últimos 6-12 meses.",
            "Baseline de OTIF, ERI y errores de picking vs. metas propuestas.",
            "Modelo financiero ROI (1 página con payback en meses).",
            "Proyección de demanda por apertura de tiendas (tabla trimestral Q1-Q4).",
            "Organigrama visual antes/después con roles y células.",
        ],
    )

    doc.add_heading("10.3 Escenarios alternativos para negociación", level=2)
    add_table(
        doc,
        ["Escenario", "Personal", "Brecha cubierta", "Ventaja", "Riesgo"],
        [
            ["A: 2 contrataciones inmediatas", "11 total", "~82% del déficit", "Menor costo", "Saturación persiste"],
            ["B: 3 contrataciones (propuesta)", "12 total", "100% + margen crecimiento", "Cubre expansión", "Mayor costo fijo"],
            ["C: 2 ahora + 1 en 6 meses", "11 → 12", "Escalonado", "Reduce riesgo financiero", "Brecha temporal en pico"],
        ],
    )

    doc.add_heading("10.4 Preguntas que Gerencia probablemente hará", level=2)
    add_table(
        doc,
        ["Pregunta", "Respuesta preparada requerida"],
        [
            ["¿Por qué 3 y no 2?", "Tabla proyección trimestral con 12,5% por tienda nueva"],
            ["¿Cuánto cuesta y cuánto recuperamos?", "ROI con payback en meses"],
            ["¿Qué pasa si contratamos 2 ahora y 1 después?", "Escenario C con cronograma"],
            ["¿Cómo medimos éxito a 90 días?", "KPIs con baseline y fecha de revisión"],
            ["¿No se puede optimizar con horas extra?", "Costo HH extra vs. costo fijo 3 personas"],
        ],
    )

    doc.add_page_break()

    # --- 11. PLAN DE ACCIÓN ---
    doc.add_heading("11. Plan de acción sugerido", level=1)
    add_table(
        doc,
        ["Paso", "Acción", "Responsable sugerido", "Plazo"],
        [
            ["1", "Aclarar nómina real (9 vs. 10) y corregir documento", "Jefe de Almacén", "Inmediato"],
            ["2", "Recopilar histórico asistencia 6-12 meses", "RRHH + Almacén", "1 semana"],
            ["3", "Medir baseline OTIF, ERI, errores picking", "Analista Inventarios", "1 semana"],
            ["4", "Time study REFRESH y conteos cíclicos", "Supervisor Almacén", "2 semanas"],
            ["5", "Construir modelo ROI con Finanzas", "Finanzas + Operaciones", "1 semana"],
            ["6", "Elaborar proyección trimestral por tiendas", "Planeación Comercial", "1 semana"],
            ["7", "Resolver flujo Consumibles entre células", "Jefe de Almacén", "1 semana"],
            ["8", "Presentar versión corregida a Gerencia", "Jefe de Almacén", "Semana 4"],
        ],
    )

    # --- 12. CONCLUSIÓN FINAL ---
    doc.add_heading("12. Conclusión final", level=1)
    doc.add_paragraph(
        "La propuesta de aumento de plantilla y reestructuración en células responde a un problema "
        "real y urgente: la operación de almacén opera al 99,6% de su capacidad real, sin margen "
        "para imprevistos, y con tareas críticas de control de inventario completamente desatendidas."
    )
    doc.add_paragraph(
        "La solución organizacional (4 células con custodia, KPIs, self-audit y protocolos de "
        "transferencia) es sólida y está alineada con buenas prácticas de gestión de almacenes."
    )
    doc.add_paragraph(
        "Sin embargo, para convertir esta propuesta de «solicitud operativa» en «caso de inversión "
        "irrebatible», es indispensable:"
    )
    add_bullet_list(
        doc,
        [
            "Corregir inconsistencias numéricas y metodológicas identificadas en este documento.",
            "Respaldar cada supuesto con datos históricos verificables.",
            "Cuantificar el ROI con cifras reales de costo, recuperación REFRESH y merma evitada.",
            "Incluir baseline de KPIs y proyección trimestral de demanda por expansión.",
            "Preparar escenarios alternativos de contratación para la negociación con Gerencia.",
        ],
    )

    p = doc.add_paragraph()
    r = p.add_run(
        "Veredicto: BUENA BASE — REQUIERE FORTALECIMIENTO ANALÍTICO ANTES DE JUNTA DIRECTIVA."
    )
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(31, 78, 121)

    # Pie de documento
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f = footer.add_run(
        "— Fin del documento —\n"
        "Revisión crítica independiente · Propuesta Team Almacén Fábrica · Agosto 2026"
    )
    f.font.size = Pt(9)
    f.italic = True
    f.font.color.rgb = RGBColor(128, 128, 128)

    return doc


if __name__ == "__main__":
    output_path = "/workspace/REVISION_CRITICA_PROPUESTA_ALMACEN.docx"
    document = build_document()
    document.save(output_path)
    print(f"Documento generado: {output_path}")
