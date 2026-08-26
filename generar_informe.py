"""Renderiza el informe de auditoria a Word (.docx) y a Markdown desde una sola fuente.

Uso:  python3 generar_informe.py
Salidas:
    /opt/cursor/artifacts/Auditoria_Propuesta_Almacen.docx
    REVISION_PROPUESTA_ALMACEN.md
"""

import os
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from informe_contenido import CONTENIDO, FECHA, SUBTITULO, TITULO

AZUL_OSCURO = RGBColor(0x1F, 0x3B, 0x60)
AZUL_MEDIO = RGBColor(0x2E, 0x6D, 0xA4)
GRIS_TEXTO = RGBColor(0x33, 0x33, 0x33)
GRIS_SUAVE = RGBColor(0x59, 0x59, 0x59)
ROJO = RGBColor(0xB3, 0x2D, 0x2D)

SOMBRA_ENCABEZADO = "1F3B60"
SOMBRA_ALTERNA = "F2F5F9"
SOMBRA_CALLOUT = "FDF6E3"
SOMBRA_CODIGO = "F4F4F4"

TOKEN = re.compile(r"(\*\*.+?\*\*|`.+?`|\*.+?\*)")


# --------------------------------------------------------------------------- utilidades docx
def _sombrear(celda, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    celda._tc.get_or_add_tcPr().append(shd)


def _sombrear_parrafo(parrafo, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    parrafo._p.get_or_add_pPr().append(shd)


def _borde_izquierdo(parrafo, hex_color, ancho=18):
    pbdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(ancho))
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), hex_color)
    pbdr.append(left)
    parrafo._p.get_or_add_pPr().append(pbdr)


def _borde_inferior(parrafo, hex_color, ancho=8):
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(ancho))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), hex_color)
    pbdr.append(bottom)
    parrafo._p.get_or_add_pPr().append(pbdr)


def _repetir_encabezado(fila):
    tr_pr = fila._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def _campo_pagina(parrafo):
    """Inserta el campo dinamico de numero de pagina."""
    run = parrafo.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    return run


ANCHO_UTIL = 6.5  # pulgadas disponibles entre margenes


def _anchos_columna(headers, rows):
    """Reparte el ancho util segun el contenido, para que la columna de etiquetas
    no quede comprimida al mismo ancho que las columnas numericas."""
    n = len(headers)
    largos = []
    for i in range(n):
        celdas = [headers[i]] + [f[i] for f in rows if i < len(f)]
        # el peso usa la raiz para amortiguar: evita que un texto muy largo se coma la tabla
        largos.append(max(len(re.sub(r"[*`]", "", c)) for c in celdas) ** 0.5)

    minimo = 0.62
    disponible = ANCHO_UTIL - minimo * n
    total = sum(largos)
    anchos = [minimo + disponible * (l / total) for l in largos]

    # tope para que ninguna columna se desborde
    for i, a in enumerate(anchos):
        if a > 3.1:
            exceso = a - 3.1
            anchos[i] = 3.1
            resto = [j for j in range(n) if j != i]
            for j in resto:
                anchos[j] += exceso / len(resto)
    return anchos


def _fijar_anchos(tabla, anchos):
    """Word y LibreOffice solo respetan los anchos si se fijan en tres sitios:
    el layout de la tabla, la rejilla (tblGrid) y cada celda."""
    tbl_pr = tabla._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)

    grid = tabla._tbl.find(qn("w:tblGrid"))
    if grid is not None:
        cols = grid.findall(qn("w:gridCol"))
        for col, ancho in zip(cols, anchos):
            col.set(qn("w:w"), str(int(ancho * 1440)))

    for fila in tabla.rows:
        for i, celda in enumerate(fila.cells):
            if i < len(anchos):
                celda.width = Inches(anchos[i])


def escribir_inline(parrafo, texto, base_size=10.5, color=GRIS_TEXTO, bold_base=False):
    """Convierte **negrita**, `mono` y *cursiva* en runs con formato."""
    for token in TOKEN.split(texto):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            run = parrafo.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`") and token.endswith("`"):
            run = parrafo.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(base_size - 0.5)
        elif token.startswith("*") and token.endswith("*") and len(token) > 2:
            run = parrafo.add_run(token[1:-1])
            run.italic = True
        else:
            run = parrafo.add_run(token)
        run.font.size = Pt(base_size)
        run.font.color.rgb = color
        if bold_base:
            run.bold = True


# --------------------------------------------------------------------------- documento
def construir_docx(ruta_salida):
    doc = Document()

    seccion = doc.sections[0]
    seccion.top_margin = Inches(0.9)
    seccion.bottom_margin = Inches(0.9)
    seccion.left_margin = Inches(1.0)
    seccion.right_margin = Inches(1.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = GRIS_TEXTO
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.12

    # ---------------- portada
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("INFORME DE AUDITORÍA TÉCNICA")
    run.font.size = Pt(11)
    run.font.color.rgb = AZUL_MEDIO
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    run = p.add_run(TITULO)
    run.font.size = Pt(23)
    run.font.color.rgb = AZUL_OSCURO
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run(SUBTITULO)
    run.font.size = Pt(12)
    run.font.color.rgb = GRIS_SUAVE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    _borde_inferior(p, "2E6DA4")
    run = p.add_run(FECHA)
    run.font.size = Pt(11)
    run.font.color.rgb = GRIS_SUAVE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(26)
    run = p.add_run(
        "Documento revisado: «Propuesta Aumento de Plantilla — Reestructuración Team "
        "Almacén Fábrica» (8 páginas)"
    )
    run.font.size = Pt(10)
    run.italic = True
    run.font.color.rgb = GRIS_SUAVE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Verificación aritmética independiente y reproducible")
    run.font.size = Pt(10)
    run.italic = True
    run.font.color.rgb = GRIS_SUAVE

    doc.add_page_break()

    # ---------------- pie de pagina con numeracion
    footer = seccion.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("Auditoría técnica — Propuesta de plantilla Almacén Fábrica    |    ")
    run.font.size = Pt(8)
    run.font.color.rgb = GRIS_SUAVE
    run_num = _campo_pagina(fp)
    run_num.font.size = Pt(8)
    run_num.font.color.rgb = GRIS_SUAVE

    # ---------------- cuerpo
    for tipo, payload in CONTENIDO:

        if tipo == "h1":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(9)
            p.paragraph_format.keep_with_next = True
            _borde_inferior(p, "2E6DA4", ancho=8)
            run = p.add_run(payload)
            run.font.size = Pt(16)
            run.font.color.rgb = AZUL_OSCURO
            run.bold = True

        elif tipo == "h2":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(5)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(payload)
            run.font.size = Pt(12.5)
            run.font.color.rgb = AZUL_MEDIO
            run.bold = True

        elif tipo == "h3":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(payload)
            run.font.size = Pt(11)
            run.font.color.rgb = GRIS_TEXTO
            run.bold = True

        elif tipo == "p":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            escribir_inline(p, payload)

        elif tipo == "bul":
            for item in payload:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_after = Pt(3)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                escribir_inline(p, item)

        elif tipo == "num":
            # numeracion manual: el estilo 'List Number' de Word comparte una sola
            # definicion de numeracion y encadenaria todas las listas del documento
            for n, item in enumerate(payload, 1):
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.left_indent = Inches(0.34)
                p.paragraph_format.first_line_indent = Inches(-0.34)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                run = p.add_run(f"{n}.\t")
                run.bold = True
                run.font.size = Pt(10.5)
                run.font.color.rgb = AZUL_MEDIO
                escribir_inline(p, item)

        elif tipo == "code":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.left_indent = Inches(0.22)
            _sombrear_parrafo(p, SOMBRA_CODIGO)
            _borde_izquierdo(p, "9AA5B1", ancho=12)
            run = p.add_run(payload)
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = GRIS_TEXTO

        elif tipo == "callout":
            tabla = doc.add_table(rows=1, cols=1)
            tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
            cant_split = OxmlElement("w:cantSplit")
            tabla.rows[0]._tr.get_or_add_trPr().append(cant_split)
            celda = tabla.cell(0, 0)
            _sombrear(celda, SOMBRA_CALLOUT)
            celda.paragraphs[0].paragraph_format.space_before = Pt(4)
            celda.paragraphs[0].paragraph_format.space_after = Pt(4)
            celda.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            escribir_inline(celda.paragraphs[0], payload, base_size=10.5, color=RGBColor(0x5C, 0x3D, 0x00))
            doc.add_paragraph().paragraph_format.space_after = Pt(2)

        elif tipo == "table":
            headers, rows, caption = payload
            tabla = doc.add_table(rows=1, cols=len(headers))
            tabla.style = "Table Grid"
            tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
            tabla.autofit = False
            anchos = _anchos_columna(headers, rows)

            hdr = tabla.rows[0]
            _repetir_encabezado(hdr)
            for i, texto in enumerate(headers):
                celda = hdr.cells[i]
                _sombrear(celda, SOMBRA_ENCABEZADO)
                par = celda.paragraphs[0]
                par.paragraph_format.space_after = Pt(2)
                par.paragraph_format.space_before = Pt(2)
                run = par.add_run(texto)
                run.bold = True
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

            for idx, fila in enumerate(rows):
                celdas = tabla.add_row().cells
                for i, texto in enumerate(fila):
                    if idx % 2 == 1:
                        _sombrear(celdas[i], SOMBRA_ALTERNA)
                    par = celdas[i].paragraphs[0]
                    par.paragraph_format.space_after = Pt(2)
                    par.paragraph_format.space_before = Pt(2)
                    escribir_inline(par, texto, base_size=9.5)

            _fijar_anchos(tabla, anchos)

            if caption:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(10)
                run = p.add_run(caption)
                run.font.size = Pt(8.5)
                run.italic = True
                run.font.color.rgb = GRIS_SUAVE
            else:
                doc.add_paragraph().paragraph_format.space_after = Pt(4)

        elif tipo == "image":
            ruta, caption = payload
            if os.path.exists(ruta):
                doc.add_picture(ruta, width=Inches(6.4))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(caption)
                run.font.size = Pt(8.5)
                run.italic = True
                run.font.color.rgb = GRIS_SUAVE

        elif tipo == "pagebreak":
            doc.add_page_break()

    doc.save(ruta_salida)
    return doc


# --------------------------------------------------------------------------- markdown
def construir_md(ruta_salida):
    lineas = [
        f"# {TITULO}",
        "",
        f"**{SUBTITULO}**",
        "",
        f"*{FECHA} — Documento revisado: «Propuesta Aumento de Plantilla — "
        "Reestructuración Team Almacén Fábrica» (8 páginas)*",
        "",
        "---",
        "",
    ]

    for tipo, payload in CONTENIDO:
        if tipo == "h1":
            lineas += [f"## {payload}", ""]
        elif tipo == "h2":
            lineas += [f"### {payload}", ""]
        elif tipo == "h3":
            lineas += [f"#### {payload}", ""]
        elif tipo == "p":
            lineas += [payload, ""]
        elif tipo == "bul":
            lineas += [f"- {i}" for i in payload] + [""]
        elif tipo == "num":
            lineas += [f"{n}. {i}" for n, i in enumerate(payload, 1)] + [""]
        elif tipo == "code":
            lineas += ["```", payload, "```", ""]
        elif tipo == "callout":
            lineas += [f"> **{payload}**", ""]
        elif tipo == "table":
            headers, rows, caption = payload
            lineas.append("| " + " | ".join(headers) + " |")
            lineas.append("|" + "|".join(["---"] * len(headers)) + "|")
            for fila in rows:
                lineas.append("| " + " | ".join(fila) + " |")
            lineas.append("")
            if caption:
                lineas += [f"*{caption}*", ""]
        elif tipo == "image":
            ruta, caption = payload
            lineas += [f"![{caption}]({ruta})", "", f"*{caption}*", ""]
        elif tipo == "pagebreak":
            lineas += ["---", ""]

    lineas += [
        "---",
        "",
        "### Reproducibilidad",
        "",
        "- `verificacion_almacen.py` — recálculo independiente de cada cifra publicada, "
        "con pruebas de consistencia interna y escenarios alternativos.",
        "- `analisis_avanzado_almacen.py` — sensibilidad cruzada, punto de saturación, "
        "atraso acumulado, nivelación de carga y umbral del REFRESH.",
        "- `grafico_carga_almacen.py` — perfil de carga por día y escenarios de dotación.",
        "- `informe_contenido.py` + `generar_informe.py` — fuente única del informe; "
        "generan la versión Word y la versión Markdown.",
        "",
    ]

    with open(ruta_salida, "w", encoding="utf-8") as fh:
        fh.write("\n".join(lineas))


if __name__ == "__main__":
    destino_docx = "/opt/cursor/artifacts/Auditoria_Propuesta_Almacen.docx"
    doc = construir_docx(destino_docx)
    construir_md("REVISION_PROPUESTA_ALMACEN.md")

    n_tablas = sum(1 for t, _ in CONTENIDO if t == "table")
    n_h1 = sum(1 for t, _ in CONTENIDO if t == "h1")
    print(f"Word generado  : {destino_docx}")
    print(f"  tamano       : {os.path.getsize(destino_docx):,} bytes")
    print(f"  parrafos     : {len(doc.paragraphs)}")
    print(f"  tablas       : {len(doc.tables)} (contenido declara {n_tablas} + callouts)")
    print(f"  capitulos    : {n_h1}")
    print("Markdown       : REVISION_PROPUESTA_ALMACEN.md")
